import sys
import json
import docx
import os
import re
import webbrowser
import urllib.parse
from openpyxl import load_workbook, Workbook
from openpyxl.styles import Font, PatternFill
import google.generativeai as genai
from dotenv import load_dotenv

# Define Fixed Headers for the Master List (Global)
HEADERS = [
    "Discretion",
    "WRAP",
    "Clients name",
    "Effective date",
    "Client signature page 11",
    "client date page 11",
    "account number",
    "Fee type",
    "Fee amount",
    "ADV received date",
    "Client signature page 14",
    "client date page 14"
]

# Extract logic - only returns JSON data
def extract_data(doc_path):
    extraction_results = {
        "status": "pending",
        "debug_notes": [],
        "gemini_data": {}
    }

    # Load env vars from .env.local explicitly
    # Script is in /scripts, .env.local is in root
    env_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), '.env.local')
    load_dotenv(env_path)

    # Initialize Gemini
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        return {"error": "GEMINI_API_KEY not found in environment", "status": "failed"}

    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('models/gemini-2.0-flash')
        
        # 1. Extract Text from DOCX
        doc_text_content = ""
        
        try:
            doc = docx.Document(doc_path)
            
            # Extract Paragraphs
            doc_text_content += "\n--- DOC CONTENT ---\n"
            for para in doc.paragraphs:
                if para.text.strip():
                   doc_text_content += para.text + "\n"
            
            # Extract Tables (often where fees are)
            doc_text_content += "\n--- TABLE CONTENT ---\n"
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        doc_text_content += " | ".join(row_text) + "\n"
                        
        except Exception as doc_err:
             extraction_results["debug_notes"].append(f"Error reading .docx: {doc_err}")
             return {"error": f"Failed to read .docx file: {doc_err}", "status": "failed"}

        if not doc_text_content:
             return {"error": "No text could be extracted from the document.", "status": "failed"}

        # 2. Query Gemini
        prompt = f"""
        Analyze the following text extracted from an Advisory Agreement and extract the specific data points.
        Return ONLY a raw JSON object. No markdown.

        CRITICAL INSTRUCTIONS:
        1. **Fee:** Look at the "Annualized Investment Advisory Fee" table. IF EMPTY, look immediately below at "Other Arrangements" or "Flat Fee". (Target value in text is often like 'Flat 1%' or similar).
        2. **Client Name:** Distinguish between the "Entity/Trust Name" (often Top line) and the "Signer Name" (often Bottom line).
        3. **Date:** Use the "Effective Date" found on Page 1.

        Data Points to Extract:

        1. "Client Entity Name" -> The legal name of the trust or entity (e.g., "Ted Smith Trust").
        2. "Authorized Signer" -> The name of the person signing (e.g., "Joshua Bajorek").
        3. "Rep Code" -> Look for a Rep Code (e.g., "TSM") usually on Page 1 or with the rep signature.
        4. "Effective Date" -> The date the agreement is effective (MM/DD/YYYY).
        5. "Account Type" -> Combine Discretionary/Non-Discretionary AND Wrap/Non-Wrap status (e.g., "Discretionary Wrap").
        6. "Fee Structure" -> The fee percentage or arrangement (e.g., "Flat 1%", "Tiered", "1.00%"). PRIORITIZE "Other Arrangements" if standard table is empty.
        7. "Account Number" -> Look for "Account Number" in registration tables.
        8. "ADV Received Date" -> Look for "date received ADV".
        9. "Client Signed P11" -> "Yes" if signature/date present on Agreement page (approx Page 11), else "No".
        10. "Client Dated P11" -> The actual date the client wrote on Page 11 (MM/DD/YYYY format), or "Not Found" if no date.
        11. "Client Signed P14" -> "Yes" if signature/date present on Fee page (approx Page 13-14), else "No".
        12. "Client Dated P14" -> The actual date the client wrote on Page 14 (MM/DD/YYYY format), or "Not Found" if no date.

        If a value is absolutely not found, return "Not Found".

        Extracted Document Text:
        {doc_text_content}
        """

        response = model.generate_content(prompt)
        
        # Clean response (remove markdown code blocks if present)
        response_text = response.text.strip()
        if response_text.startswith("```json"):
            response_text = response_text[7:]
        if response_text.endswith("```"):
            response_text = response_text[:-3]
            
        data = json.loads(response_text)
        extraction_results["gemini_data"] = data
        
        # Backward compatibility fee key (though less relevant now)
        extraction_results["fee"] = data.get('Fee Structure', 'Not Found')
        extraction_results["status"] = "success"
        
        # Include document text for frontend preview
        extraction_results["document_text"] = doc_text_content

        return extraction_results

    except Exception as e:
        return {"error": str(e), "status": "failed", "debug_notes": extraction_results["debug_notes"]}


# Save logic - Updates Excel and Email
def save_data(data, excel_path):
    save_results = {
        "status": "pending",
        "debug_notes": []
    }

    try:
        # 3. Update Excel
        os.makedirs(os.path.dirname(excel_path), exist_ok=True)
        
        # Open or Create Workbook
        if os.path.exists(excel_path):
             try:
                 wb = load_workbook(excel_path)
             except:
                 wb = Workbook()
        else:
             wb = Workbook()
             
        ws = wb.active
        
        # Define Fixed Headers for the Master List
        HEADERS = [
            "Discretion",
            "WRAP",
            "Clients name",
            "Effective date",
            "Client signature page 11",
            "client date page 11",
            "account number",
            "Fee type",
            "Fee amount",
            "ADV received date",
            "Client signature page 14",
            "client date page 14"
        ]
        
        # Check if headers exist (row 1), if not, write them
        if ws.max_row == 0 or (ws.max_row == 1 and ws['A1'].value is None):
            for col_num, header in enumerate(HEADERS, 1):
                cell = ws.cell(row=1, column=col_num, value=header)
                cell.font = Font(bold=True)
        
        # Prepare row data
        next_row = ws.max_row + 1
        
        # Map data keys to our Ordered Headers
        for col_num, header in enumerate(HEADERS, 1):
            val = data.get(header, "") 
            ws.cell(row=next_row, column=col_num, value=val)
            
        # Adjust column widths for readability
        for col in range(1, len(HEADERS) + 1):
             ws.column_dimensions[chr(64 + col)].width = 20

        wb.save(excel_path)
        
        save_results["status"] = "success"

        # 4. Open Outlook (Mailto)
        try:
            subject = "AA Updated - Analysis Complete"
            
            # Format email body with all extracted data points
            # Use direct keys sent from frontend
            discretionary = data.get('Discretion', 'Not Found')
            wrap = data.get('WRAP', 'Not Found')
            
            fee_type = data.get('Fee type', 'Not Found')
            fee_amount = data.get('Fee amount', 'Not Found')
            
            body_lines = [
                "Advisory Agent Analysis Results:",
                "",
                f"Discretionary v. Non-Discretionary: {discretionary}",
                f"Wrap v. Non-WRAP: {wrap}",
                f"Client's Name: {data.get('Clients name', 'Not Found')}",
                f"Effective Date: {data.get('Effective date', 'Not Found')}",
                f"Client Signed Page 11: {data.get('Client signature page 11', 'Not Found')}",
                f"Client Dated Page 11: {data.get('client date page 11', 'Not Found')}",
                f"Account Number: {data.get('account number', 'Not Found')}",
                f"Flat v. Tiered: {fee_type}",
                f"Fee Amount: {fee_amount}",
                f"ADV Received Date: {data.get('ADV received date', 'Not Found')}",
                f"Client Signed Page 14: {data.get('Client signature page 14', 'Not Found')}",
                f"Client Dated Page 14: {data.get('client date page 14', 'Not Found')}",
            ]
            
            body_text = "\n".join(body_lines)
            
            params = {
                "subject": subject,
                "body": body_text
            }
            query_string = urllib.parse.urlencode(params).replace("+", "%20")
            mailto_link = f"mailto:?{query_string}"
            
            webbrowser.open(mailto_link)
            
        except Exception as mail_err:
             save_results["debug_notes"].append(f"Failed to open mail client: {mail_err}")

        return save_results

    except Exception as e:
        return {"error": str(e), "status": "failed", "debug_notes": save_results["debug_notes"]}

def send_nigo_email(data):
    nigo_results = {
        "status": "pending",
        "debug_notes": []
    }
    
    try:
        # --- 1. Update NIGO Excel Sheet ---
        excel_path = r"C:\Users\JoshuaBajorek\Desktop\Advisory Agent.xlsx"
        try:
             # Open or Create Workbook
            if os.path.exists(excel_path):
                 try:
                     wb = load_workbook(excel_path)
                 except:
                     wb = Workbook()
            else:
                 wb = Workbook()
            
            # Get or Create 'NIGO' sheet
            if "NIGO" in wb.sheetnames:
                ws = wb["NIGO"]
            else:
                ws = wb.create_sheet("NIGO")
            
            # Check Headers
            if ws.max_row == 0 or (ws.max_row == 1 and ws['A1'].value is None):
                for col_num, header in enumerate(HEADERS, 1):
                    cell = ws.cell(row=1, column=col_num, value=header)
                    cell.font = Font(bold=True)
            
            # Prepare row data
            next_row = ws.max_row + 1
            
            # Red Fill for invalid cells
            red_fill = PatternFill(start_color="FFCCCC", end_color="FFCCCC", fill_type="solid")
            
            # Map data keys and write
            for col_num, header in enumerate(HEADERS, 1):
                val = data.get(header, "")
                cell = ws.cell(row=next_row, column=col_num, value=val)
                
                # Check validity for highlighting
                # Invalid if empty, "Not Found", "Missing", "Error", or "No" for specific yes/no fields
                is_invalid = False
                val_str = str(val).lower() if val else ""
                
                if not val or val_str in ["not found", "missing", "error"]:
                    is_invalid = True
                elif header in ["Client signature page 11", "client date page 11", "Client signature page 14", "client date page 14"]:
                    if val_str != "yes":
                        is_invalid = True
                
                if is_invalid:
                    cell.fill = red_fill
            
            # Adjust widths
            for col in range(1, len(HEADERS) + 1):
                 ws.column_dimensions[chr(64 + col)].width = 20

            wb.save(excel_path)
            nigo_results["excel_updated"] = True
            
        except Exception as excel_err:
            nigo_results["debug_notes"].append(f"Failed to update NIGO Excel: {excel_err}")

        # --- 2. Open Outlook (Mailto) ---
        subject = "Advisory Agreement NIGO"
        
        # Identify missing items
        missing_items = []
        
        # Check specific entries
        # Use new keys from headers
        checks = [
            ("Client Entity Name", data.get("Clients name", "")),
            ("Effective Date", data.get("Effective date", "")),
            ("Client Signed P11", data.get("Client signature page 11", "")),
            ("Client Dated P11", data.get("client date page 11", "")),
            ("Account Number", data.get("account number", "")),
            ("ADV Received Date", data.get("ADV received date", "")),
            ("Client Signed P14", data.get("Client signature page 14", "")),
            ("Client Dated P14", data.get("client date page 14", "")),
            ("Fee Amount", data.get("Fee amount", "")),
            ("Fee Type", data.get("Fee type", "")),
            ("Discretionary Status", data.get("Discretion", "")),
            ("Wrap Status", data.get("WRAP", ""))
        ]
        
        for label, val in checks:
            # Check for empty, "Not Found", "Missing", "Error"
            if not val or val in ["Not Found", "Missing", "Error"]:
                missing_items.append(label)
        
        body_lines = [
            "Please see the following missing items that need to be corrected:",
            ""
        ]
        
        if missing_items:
            for item in missing_items:
                body_lines.append(f"- {item}")
        else:
            body_lines.append("No specific extracted fields marked as missing/error. Please review document manually.")
            
        body_lines.append("")
        body_lines.append("Please resubmit once corrected.")

        body_text = "\n".join(body_lines)
        
        params = {
            "subject": subject,
            "body": body_text
        }
        query_string = urllib.parse.urlencode(params).replace("+", "%20")
        mailto_link = f"mailto:?{query_string}"
        
        webbrowser.open(mailto_link)
        nigo_results["status"] = "success"
        return nigo_results

    except Exception as e:
        return {"error": str(e), "status": "failed", "debug_notes": nigo_results["debug_notes"]}

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print(json.dumps({"error": "Missing arguments. Usage: python selenium_tool.py [extract|save] [doc_path|json_data]"}))
        sys.exit(1)
        
    mode = sys.argv[1]
    payload = sys.argv[2]
    
    # Hardcoded path as requested
    excel_path = r"C:\Users\JoshuaBajorek\Desktop\Advisory Agent.xlsx"
    
    if mode == "extract":
        result = extract_data(payload)
        print(json.dumps(result))
    elif mode == "save":
        try:
            # When calling from Node, JSON might be escaped, so parse it
            data = json.loads(payload)
            result = save_data(data, excel_path)
            print(json.dumps(result))
        except json.JSONDecodeError as e:
             print(json.dumps({"error": f"Invalid JSON payload for save mode: {e}", "payload_snippet": payload[:100]}))
    elif mode == "nigo":
        try:
            data = json.loads(payload)
            result = send_nigo_email(data)
            print(json.dumps(result))
        except json.JSONDecodeError as e:
             print(json.dumps({"error": f"Invalid JSON payload for nigo mode: {e}", "payload_snippet": payload[:100]}))
    else:
        print(json.dumps({"error": f"Unknown mode: {mode}"}))
